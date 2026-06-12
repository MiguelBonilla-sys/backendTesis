"""
Conversor Markdown → DOCX con formato APA institucional USB.

Toma el documento maestro en Markdown (docs/TESIS-FINAL.md) y produce un .docx
que hereda los estilos de la plantilla oficial "Plantilla APA Tesis 2022 v6.docx"
(Times New Roman 12, justificado, interlineado 1.5, márgenes 2.5 cm, Heading 1
centrado en negrita, etc.). El resultado queda listo para entregar tras pegar la
portada institucional y las imágenes de las figuras.

Diseño: en lugar de reconstruir los estilos a mano, abre la plantilla como base
(así conserva definiciones de estilo, márgenes y numeración), vacía su cuerpo de
texto de ejemplo y vuelca el contenido del Markdown mapeando cada elemento al
estilo APA correspondiente.

Uso:
    python -m scripts.md_to_docx_apa \
        --md ../docs/TESIS-FINAL.md \
        --template "/ruta/Plantilla APA Tesis 2022 v6.docx" \
        --out ../docs/TESIS-FINAL.docx

Requiere: python-docx (uv pip install python-docx)

Mapeo Markdown → estilo APA:
    # H1            → Title (portada/título principal)
    ## H1 numérico  → Heading 1 (capítulo, centrado negrita)
    ### H2          → Heading 2 (subcapítulo, izquierda negrita)
    #### H3         → Heading 3
    párrafo         → Párr.APA si existe, si no Normal (sangría 1ª línea)
    > cita          → Cita+40 (cita en bloque ≥ 40 palabras)
    - viñeta        → List Bullet
    | tabla |       → tabla con estilo Table Grid + Caption
    ```código```    → Normal monoespaciado (Consolas 10)
    **negrita**, `code` inline → runs con formato
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

try:
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor
except ImportError:
    print("Falta python-docx. Instalá con: uv pip install python-docx", file=sys.stderr)
    sys.exit(2)


# --------------------------------------------------------------------------- #
# Utilidades de estilo
# --------------------------------------------------------------------------- #

def _style_exists(doc, name: str) -> bool:
    try:
        _ = doc.styles[name]
        return True
    except KeyError:
        return False


def _clear_body(doc) -> None:
    """Elimina todos los párrafos y tablas del cuerpo de la plantilla."""
    body = doc.element.body
    for child in list(body):
        # conservar la última <w:sectPr> (configuración de página/márgenes)
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def _add_inline_runs(paragraph, text: str) -> None:
    """Agrega runs interpretando **negrita** y `code` inline."""
    # tokeniza en segmentos de negrita, código y texto plano
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    for token in pattern.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(token)


# --------------------------------------------------------------------------- #
# Parser de Markdown (línea por línea, sin dependencias extra)
# --------------------------------------------------------------------------- #

def _emit_table(doc, rows: list[list[str]]) -> None:
    """Vuelca una tabla Markdown como tabla DOCX (Table Grid)."""
    if len(rows) < 2:
        return
    header = rows[0]
    body_rows = rows[2:]  # rows[1] es el separador |---|
    table = doc.add_table(rows=1, cols=len(header))
    if _style_exists(doc, "Table Grid"):
        table.style = "Table Grid"
    for i, cell_text in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].text = ""
        _add_inline_runs(cell.paragraphs[0], cell_text.strip())
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in body_rows:
        cells = table.add_row().cells
        for i in range(len(header)):
            txt = row[i].strip() if i < len(row) else ""
            cells[i].paragraphs[0].text = ""
            _add_inline_runs(cells[i].paragraphs[0], txt)


def _split_table_row(line: str) -> list[str]:
    parts = line.strip().strip("|").split("|")
    return [p.strip() for p in parts]


def convert(md_path: Path, template_path: Path, out_path: Path) -> None:
    doc = docx.Document(str(template_path))
    _clear_body(doc)

    para_style = "Párr.APA" if _style_exists(doc, "Párr.APA") else "Normal"
    quote_style = "Cita+40" if _style_exists(doc, "Cita+40") else "Normal"

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_buf: list[str] = []

    while i < len(lines):
        line = lines[i]

        # bloque de código cercado ```
        if line.strip().startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Consolas"
                run.font.size = Pt(10)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        stripped = line.strip()

        # separador horizontal --- → salto, se ignora como contenido
        if stripped == "---":
            i += 1
            continue

        # encabezados
        if stripped.startswith("#"):
            m = re.match(r"(#{1,6})\s+(.*)", stripped)
            if m:
                level = len(m.group(1))
                text = m.group(2).strip()
                if level == 1:
                    style = "Title" if _style_exists(doc, "Title") else "Heading 1"
                elif level == 2:
                    style = "Heading 1"
                elif level == 3:
                    style = "Heading 2"
                else:
                    style = "Heading 3"
                p = doc.add_paragraph(style=style)
                _add_inline_runs(p, text)
                i += 1
                continue

        # cita en bloque >
        if stripped.startswith(">"):
            text = re.sub(r"^>\s?", "", stripped)
            p = doc.add_paragraph(style=quote_style)
            _add_inline_runs(p, text)
            i += 1
            continue

        # tabla Markdown
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(
            r"^\|[\s:|-]+\|?$", lines[i + 1].strip()
        ):
            tbl_rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl_rows.append(_split_table_row(lines[i]))
                i += 1
            _emit_table(doc, tbl_rows)
            continue

        # viñetas
        if re.match(r"^[-*]\s+", stripped):
            text = re.sub(r"^[-*]\s+", "", stripped)
            bullet_style = "List Bullet" if _style_exists(doc, "List Bullet") else para_style
            p = doc.add_paragraph(style=bullet_style)
            _add_inline_runs(p, text)
            i += 1
            continue

        # línea en blanco
        if not stripped:
            i += 1
            continue

        # párrafo normal
        p = doc.add_paragraph(style=para_style)
        _add_inline_runs(p, stripped)
        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"DOCX generado → {out_path}")
    print("Pasos manuales restantes: pegar portada institucional, insertar las "
          "imágenes de las figuras y actualizar la tabla de contenido en Word "
          "(References → Update Table).")


def main() -> int:
    here = Path(__file__).resolve().parent
    default_md = here.parent.parent / "docs" / "TESIS-FINAL.md"
    default_out = here.parent.parent / "docs" / "TESIS-FINAL.docx"
    default_tpl = Path.home() / "Downloads" / "Plantilla APA Tesis 2022 v6.docx"

    parser = argparse.ArgumentParser(description="Markdown → DOCX (formato APA USB)")
    parser.add_argument("--md", default=str(default_md), help="Markdown de entrada")
    parser.add_argument("--template", default=str(default_tpl), help="Plantilla .docx APA")
    parser.add_argument("--out", default=str(default_out), help="Salida .docx")
    args = parser.parse_args()

    md_path = Path(args.md)
    tpl_path = Path(args.template)
    if not md_path.exists():
        print(f"No existe el Markdown: {md_path}", file=sys.stderr)
        return 1
    if not tpl_path.exists():
        print(f"No existe la plantilla: {tpl_path}", file=sys.stderr)
        return 1

    convert(md_path, tpl_path, Path(args.out))
    return 0


if __name__ == "__main__":
    sys.exit(main())
